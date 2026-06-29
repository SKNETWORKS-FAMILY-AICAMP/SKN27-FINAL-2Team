from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = max(Path(r"C:\Users\Playdata\Downloads").glob("*docx*"), key=lambda p: p.stat().st_mtime)
OUTPUT = ROOT / "docs" / "데이터_전처리_결과서_27기_2팀.docx"


def set_cell(cell, text: str) -> None:
    cell.text = text


def resize_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)

    for row, values in zip(table.rows, rows):
        while len(values) < len(row.cells):
            values.append("")
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def replace_paragraphs(doc: Document) -> None:
    replacements = {
        "전처리 목적 : 멀티 에이전트·RAG 파이프라인에 투입 가능한 학습·검색용 고품질 데이터 생산": (
            "전처리 목적 : 한국사 학습 챗봇과 RAG 검색에 투입 가능한 고품질 문서·청크·연표 데이터 생산"
        ),
        "전처리 범위 : 정형 DB 데이터 및 비정형 문서·로그(원천 → 정제 → 청킹·임베딩)": (
            "전처리 범위 : 사료로 본 한국사, 신편 한국사, 한국사 이미지 자료, 한국사 연대기 연표"
            "(원천 → 정제 → 메타데이터 생성 → 청킹·임베딩·DB 적재)"
        ),
        "사용 도구 : pandas, numpy, re, kiwipiepy, nltk, scikit-learn, LangChain, sentence-transformers": (
            "사용 도구 : Python, csv, json, re, pathlib, psycopg2, PostgreSQL, pgvector, OpenAI Embedding API"
        ),
        "정제 · 가공 · 분할 절차 및 후속 활용": "정제 · 메타데이터 생성 · 청크 분할 · DB 적재 절차 및 후속 활용",
        "분할 기준 : 층화 추출 (category 기준, random_state=42)": (
            "분할 기준 : 학습/검증/테스트 분할 대신 RAG 검색 단위 문서·청크 및 연표 보조 테이블로 분리"
        ),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]


def main() -> None:
    doc = Document(str(TEMPLATE))
    replace_paragraphs(doc)

    resize_table(
        doc.tables[1],
        [
            ["산출물 단계", "데이터 전처리"],
            ["제출 일자", "2026. 06. 29."],
            ["깃허브 경로", "SKN27-FINAL-2Team"],
            ["작성 팀원", "2팀"],
        ],
    )

    resize_table(
        doc.tables[3],
        [
            ["이슈 유형", "대상 필드", "판정 기준", "처리 방식", "처리량 / 비고"],
            ["결측치", "title / content / description", "빈 문자열, NULL, 공백만 존재", "유효 제목·본문이 있는 행만 문서화", "총 문서 9,005건 생성 기준"],
            ["중복/식별", "자료ID / 페이지ID / 이미지ID / doc_id / chunk_id", "동일 ID 또는 원천 식별자 누락", "고유 doc_id, document_id, chunk_id 생성", "문서 9,005건 / 청크 33,821건 추적"],
            ["본문 정제", "content / 본문 / 국문 / 해설 / description", "중복 공백, 과도한 줄바꿈, 제어문자 포함", "공백·줄바꿈 정규화 및 본문 정리", "전체 문서 9,005건 대상"],
            ["메타데이터", "period / field / category / keywords / metadata", "검색 필터 및 출처 추적에 필요한 값", "JSON metadata 생성 및 보존", "문서 9,005건 metadata 구성"],
            ["청크 분할", "content / chunk_text", "RAG 검색 단위로 분할 가능한 본문", "문단 기반 청크 분할", "총 33,821건"],
            ["이미지 자료", "이미지ID / 설명 / 썸네일URL / 원본이미지URL", "이미지 검색 응답에 필요한 설명·URL 존재", "이미지 설명과 URL을 metadata에 저장", "문서 1,417건 / 청크 1,417건"],
            ["연표 데이터", "content_type / year / title / era / field", "시대 코드 + 분야 코드 조합", "content_type을 era, field로 변환", "1,162건"],
            ["불필요 컬럼", "content_type / level_id / source_url", "최종 연표 조회에 직접 사용하지 않음", "정제 CSV에서 컬럼 제거", "1,162건"],
        ],
    )

    resize_table(
        doc.tables[4],
        [
            ["처리 단계", "입력", "출력", "목적"],
            ["형식 정합성", "원천 CSV / Markdown", "UTF-8 정제 텍스트", "인코딩·스키마 통일"],
            ["본문 정제", "제목·본문·설명·해설", "공백·줄바꿈 정규화 텍스트", "임베딩 입력 안정화"],
            ["메타데이터 생성", "시대·분야·유형·키워드", "period / field / category / metadata", "검색 필터 및 출처 추적"],
            ["분할·청킹", "정제 문서", "documents.jsonl + chunks.jsonl", "RAG 검색 단위 구성"],
            ["DB 적재", "processed JSONL / CSV", "rag.document_chunks + rag.history_timeline", "챗봇 RAG 조회 활용"],
        ],
    )

    resize_table(
        doc.tables[5],
        [
            ["필드", "타입", "결측 처리", "중복 처리", "이상치 처리", "비고"],
            ["doc_id", "STRING", "원천 ID 없을 시 안정 ID 생성", "고유 ID 유지", "-", "문서 식별자"],
            ["chunk_id", "STRING", "필수", "고유 청크 ID 유지", "-", "검색 결과 추적"],
            ["content / chunk_text", "TEXT", "빈 본문 제외", "동일 문서 중복 최소화", "긴 본문 청크 분할", "임베딩 대상"],
            ["period / era", "CATEGORY", "가능한 경우 원천값·코드에서 생성", "-", "-", "시대 필터"],
            ["field", "CATEGORY", "content_type 코드에서 생성", "-", "-", "분야 필터"],
            ["metadata", "JSONB", "없으면 빈 객체", "-", "-", "출처·키워드·이미지URL 보존"],
        ],
    )

    resize_table(
        doc.tables[6],
        [
            ["구분", "항목", "기준 / 결과", "설명 및 관리 방안"],
            ["산출물", "사료 문서", "1,146건", "사료로 본 한국사 문서 단위 데이터"],
            ["산출물", "사료 청크", "7,540건", "사료 RAG 검색 단위"],
            ["산출물", "신편 한국사 문서", "6,442건", "개념·흐름 설명용 문서"],
            ["산출물", "신편 한국사 청크", "24,864건", "장문 본문을 검색 단위로 분할"],
            ["산출물", "이미지 자료 문서/청크", "각 1,417건", "이미지·유물·유적 검색용 자료"],
            ["산출물", "연표 데이터", "1,162건", "시대·분야·연도 기반 보조 조회"],
            ["품질", "전체 문서", "9,005건", "RAG 문서 단위 JSONL"],
            ["품질", "전체 청크", "33,821건", "pgvector 임베딩 및 키워드 검색 대상"],
            ["후속 활용", "RAG 검색", "rag.document_chunks", "벡터 검색 + 키워드 검색"],
            ["후속 활용", "연표 조회", "rag.history_timeline", "연표·순서·흐름 질문 보조"],
        ],
    )

    resize_table(
        doc.tables[7],
        [
            ["변경일", "변경자", "변경내용", "영향 받는 항목", "비고"],
            ["2026.06.29", "2팀", "전체 데이터 전처리 결과 정리", "t3, t4, t5, t6", "v1.0"],
            ["2026.06.29", "2팀", "연표 코드 변환 및 rag.history_timeline 적재 반영", "t3, t6", "v1.1"],
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
